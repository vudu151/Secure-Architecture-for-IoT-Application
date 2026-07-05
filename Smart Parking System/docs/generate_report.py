import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set explicit padding for table cells in dxa (1 pt = 20 dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block(doc, code_text):
    """Add a professional looking code block with gray background and Consolas font"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.0)
    run.font.color.rgb = RGBColor(40, 40, 40)
    
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    pPr.append(shd)
    
    # Add top and bottom light borders
    pbdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5 pt
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'E0E0E0')
        pbdr.append(border)
    pPr.append(pbdr)

def add_styled_table(doc, headers, data_rows, col_widths=None):
    """Helper to create a beautiful, HUST-standard styled table with alternate row colors"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1A237E")  # Dark Blue header
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        
        # Set header font style
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                
    for row_idx, row_data in enumerate(data_rows):
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(row_data):
            row_cells[col_idx].text = str(text)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
            
            # Formatting cells
            for paragraph in row_cells[col_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10.5)
            
            # Alternate row backgrounds
            if row_idx % 2 == 1:
                set_cell_background(row_cells[col_idx], "F0F2F5")
                
    # Apply column widths if specified
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)

def create_document():
    doc = Document()

    # --- Setup Page Margins (HUST Standard: Top/Bottom 2cm, Left 3cm, Right 2cm) ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.79)      # 2.0 cm
        section.bottom_margin = Inches(0.79)   # 2.0 cm
        section.left_margin = Inches(1.18)     # 3.0 cm
        section.right_margin = Inches(0.79)    # 2.0 cm

    # --- Setup Base Font (Times New Roman) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(6)

    # ==================== PAGE 1: COVER PAGE ====================
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_univ = p_univ.add_run("ĐẠI HỌC BÁCH KHOA HÀ NỘI\nTRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG\n------------------***------------------")
    run_univ.bold = True
    run_univ.font.size = Pt(14)

    # Spacing
    for _ in range(5):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("BÁO CÁO ĐỀ TÀI TIỂU LUẬN\n\nXÂY DỰNG KIẾN TRÚC BẢO MẬT TOÀN DIỆN\nCHO HỆ THỐNG IoT SMART PARKING")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(128, 0, 0) # Dark red accent

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("\n\nChuyên đề: Kiến trúc Bảo mật cho Ứng dụng IoT\nChương trình Đào tạo Thạc sĩ CNTT")
    run_sub.italic = True
    run_sub.font.size = Pt(13)

    for _ in range(6):
        doc.add_paragraph()

    # Student / Teacher Info Box
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_info.paragraph_format.left_indent = Inches(1.5)
    
    r_info = p_info.add_run(
        "Học viên thực hiện : Nguyễn Văn A\n"
        "Mã số học viên     : M24XXXXXX\n"
        "Lớp                : Cao học CNTT - Khóa 2024\n"
        "Giảng viên hướng dẫn: PGS. TS. Nguyễn Văn B"
    )
    r_info.font.size = Pt(13)
    r_info.bold = True

    for _ in range(4):
        doc.add_paragraph()

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_footer.add_run("HÀ NỘI, NĂM 2026")
    r_foot.bold = True
    r_foot.font.size = Pt(12)

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS PLACEHOLDER ====================
    h_toc = doc.add_heading(level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = h_toc.add_run("MỤC LỤC")
    r_toc.font.size = Pt(16)
    r_toc.bold = True
    r_toc.font.color.rgb = RGBColor(0, 0, 0)

    p_toc_desc = doc.add_paragraph()
    p_toc_desc.add_run(
        "CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI & KIẾN TRÚC TỔNG QUAN.......................................... 3\n"
        "1.1. Đặt vấn đề và Mục tiêu đề tài....................................................................................................... 3\n"
        "1.2. Sơ đồ kiến trúc tổng quan hệ thống................................................................................................. 3\n"
        "1.3. Mô tả các thành phần và công nghệ tích hợp.............................................................................. 4\n"
        "CHƯƠNG 2: THIẾT KẾ CƠ SỞ DỮ LIỆU & CORE BACKEND.......................................... 6\n"
        "2.1. Thực thể cơ sở dữ liệu chi tiết (Database Schema)......................................................................... 6\n"
        "2.2. Xây dựng tầng REST APIs nghiệp vụ (Spring Boot)...................................................................... 9\n"
        "2.3. Tích hợp kết nối WebSocket STOMP real-time............................................................................ 12\n"
        "2.4. Hướng dẫn cài đặt và kết nối Cơ sở dữ liệu (PostgreSQL).................................................................... 14\n"
        "CHƯƠNG 3: GIẢI PHÁP BẢO MẬT KIẾN TRÚC HỆ THỐNG........................................... 15\n"
        "3.1. Xác thực phân quyền không trạng thái (JWT Stateless Auth & Blacklist)........................................ 15\n"
        "3.2. Mã hóa dữ liệu lưu trữ (AES-256 License Plate Encryption)....................................................... 17\n"
        "3.3. Cơ chế chống tấn công phát lại (Redis Nonce Anti-Replay)......................................................... 18\n"
        "3.4. Kiểm toán và ghi nhật ký bảo mật (Security Audit Logging & Spring AOP)................................. 20\n"
        "3.5. Giới hạn tần suất yêu cầu (Rate Limiting với Bucket4j).................................................................. 22\n"
        "3.6. Cấu hình cổng bảo mật API Gateway và SSL Termination (Nginx)............................................... 23\n"
        "CHƯƠNG 4: LẬP TRÌNH GIAO DIỆN CLIENT & THIẾT BỊ IoT....................................... 26\n"
        "4.1. Cấu trúc tổ chức mã nguồn Monorepo.......................................................................................... 26\n"
        "4.2. Giao diện & Xử lý Ứng dụng di động tài xế (Flutter Mobile Client)............................................... 27\n"
        "4.3. Trang điều khiển giám sát trung tâm (ReactJS Admin Dashboard)................................................ 30\n"
        "4.4. Lập trình thiết bị IoT, Barrier & Hạ tầng mTLS Mosquitto........................................................... 32\n"
        "KẾT LUẬN & HƯỚNG PHÁT TRIỂN................................................................................. 35"
    )
    p_toc_desc.style.paragraph_format.line_spacing = 1.4

    doc.add_page_break()

    # ==================== CHƯƠNG 1 ====================
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("CHƯƠNG 1: GIỚI THIỆU ĐỀ TÀI & KIẾN TRÚC TỔNG QUAN")
    r1.font.size = Pt(15)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0, 0, 0)

    # 1.1
    h1_1 = doc.add_heading(level=2)
    r1_1 = h1_1.add_run("1.1. Đặt vấn đề và Mục tiêu đề tài")
    r1_1.font.size = Pt(13)
    r1_1.bold = True
    r1_1.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Trong bối cảnh đô thị hóa nhanh chóng, nhu cầu về đỗ xe tại các trung tâm thương mại, tòa nhà thông minh ngày càng tăng cao. "
        "Các hệ thống quản lý bãi đỗ xe truyền thống thường đối mặt với các vấn đề về tắc nghẽn giao thông tại cổng ra vào, khó khăn trong tìm kiếm vị trí trống, và rủi ro mất mát an ninh thông tin. "
        "Kiến trúc của các hệ thống IoT hiện nay cũng đặt ra nhiều lo ngại về bảo mật khi dữ liệu truyền qua môi trường internet không dây dễ bị tấn công nghe lén (Eavesdropping), tấn công phát lại (Replay attack), hay giả mạo thiết bị (Device Spoofing)."
    )
    doc.add_paragraph(
        "Đề tài này hướng đến việc thiết kế và xây dựng một hệ thống IoT Smart Parking hoàn chỉnh, áp dụng các kỹ thuật bảo mật kiến trúc phân lớp từ phần cứng lên phần mềm bao gồm:\n"
        "- Xác thực và mã hóa đường truyền qua giao thức mTLS (Mutual TLS) bảo vệ kết nối giữa thiết bị IoT (ESP32, Raspberry Pi) và Broker MQTT.\n"
        "- Bảo mật dữ liệu nhạy cảm của người dùng (biển số xe) lưu trữ trong cơ sở dữ liệu bằng thuật toán mã hóa đối xứng AES-256.\n"
        "- Phòng chống tấn công phát lại bằng cơ chế Nonce & Timestamp lưu giữ trên bộ nhớ đệm phân tán Redis.\n"
        "- Giới hạn tần suất yêu cầu chống brute-force đăng nhập sử dụng giải thuật Token Bucket thông qua thư viện Bucket4j.\n"
        "- Ghi nhận nhật ký kiểm toán bảo mật (Security Audit Logging) tự động qua AOP để phục vụ công tác giám sát điều tra sự cố."
    )

    # 1.2
    h1_2 = doc.add_heading(level=2)
    r1_2 = h1_2.add_run("1.2. Sơ đồ kiến trúc tổng quan hệ thống")
    r1_2.font.size = Pt(13)
    r1_2.bold = True
    r1_2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Hệ thống Smart Parking được xây dựng trên một kiến trúc phân lớp hướng dịch vụ (Service-Oriented Architecture - SOA) tích hợp công nghệ IoT. "
        "Sự trao đổi thông tin diễn ra thông qua ba phương thức giao tiếp chính: RESTful API, WebSocket STOMP và MQTT qua TLS. Dưới đây là mô tả luồng giao tiếp dữ liệu:"
    )

    # Chèn ảnh sơ đồ kiến trúc
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(6)
    r_img = p_img.add_run()
    r_img.add_picture('docs/architecture_diagram.png', width=Inches(5.8))
    
    # Chú thích ảnh
    p_caption = doc.add_paragraph()
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_before = Pt(2)
    p_caption.paragraph_format.space_after = Pt(12)
    r_caption = p_caption.add_run("Hình 1.1: Sơ đồ kiến trúc tổng quan và bảo mật hệ thống IoT Smart Parking")
    r_caption.italic = True
    r_caption.bold = True
    r_caption.font.size = Pt(10.5)

    # Communication flow list
    p_layer1 = doc.add_paragraph()
    p_layer1.add_run("1. Lớp Thiết bị IoT (Edge Devices Layer):").bold = True
    doc.add_paragraph("   - ESP32 chịu trách nhiệm đọc cảm biến siêu âm HC-SR04 định kỳ 1 giây để cập nhật trạng thái trống/bận tại các vị trí đỗ. ESP32 kết nối tới Mosquitto Broker trên cổng 8883 thông qua giao thức mTLS sử dụng chứng chỉ X.509 riêng biệt được cấp phát.", style='List Bullet')
    doc.add_paragraph("   - Raspberry Pi quản lý camera IP đặt tại các cổng kiểm soát ra vào. Khi có xe tới cổng, RPi chạy thuật toán nhận dạng biển số OCR (YOLOv8 Nano) hoặc quét mã QR trên vé đỗ xe của tài xế. Dữ liệu biển số xe nhận diện được gửi về Backend thông qua HTTP REST API bảo mật.", style='List Bullet')
    
    p_layer2 = doc.add_paragraph()
    p_layer2.add_run("2. Lớp Hạ tầng và Trung gian (Broker & Gateway Layer):").bold = True
    doc.add_paragraph("   - Mosquitto Broker chạy độc lập trong container, áp dụng phân quyền ACL (Access Control List) gắt gao. Chỉ thiết bị có Certificate Common Name (CN) tương ứng mới được phép viết (publish) lên topic trạng thái và đọc (subscribe) topic điều khiển barie cổng.", style='List Bullet')
    doc.add_paragraph("   - Nginx làm API Gateway và thực hiện TLS termination. Mọi yêu cầu từ Web Admin Dashboard và Mobile App đều phải đi qua Nginx cổng 443 trước khi định tuyến nội bộ vào Spring Boot App cổng 8080.", style='List Bullet')

    p_layer3 = doc.add_paragraph()
    p_layer3.add_run("3. Lớp Xử lý Trung tâm (Backend Core Layer):").bold = True
    doc.add_paragraph("   - Spring Boot Backend chịu trách nhiệm xử lý logic nghiệp vụ đặt chỗ (Booking), quản lý ví điện tử của tài xế, quản lý đăng ký phương tiện và điều khiển từ xa thiết bị. Backend duy trì kết nối WebSocket STOMP tới Web Admin Dashboard để cập nhật trạng thái sơ đồ bãi đỗ thời gian thực.", style='List Bullet')

    # ==================== CHƯƠNG 2 ====================
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("CHƯƠNG 2: THIẾT KẾ CƠ SỞ DỮ LIỆU & CORE BACKEND")
    r2.font.size = Pt(15)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0, 0, 0)

    # 2.1
    h2_1 = doc.add_heading(level=2)
    r2_1 = h2_1.add_run("2.1. Thực thể cơ sở dữ liệu chi tiết (Database Schema)")
    r2_1.font.size = Pt(13)
    r2_1.bold = True
    r2_1.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Cơ sở dữ liệu được thiết kế trên hệ quản trị cơ sở dữ liệu quan hệ PostgreSQL 16. "
        "Hệ thống quản lý thông tin thông qua 7 thực thể chính. Cấu trúc thiết kế chi tiết các trường, kiểu dữ liệu, các khóa liên kết ngoại và các ràng buộc được thể hiện cụ thể ở các bảng dưới đây:"
    )

    # Table 1: Users
    p_tbl1 = doc.add_paragraph()
    p_tbl1.add_run("Bảng 2.1: Cấu trúc thực thể users (Người dùng)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "email", "VARCHAR(255)", "Unique", "NOT NULL", "Email đăng nhập, định dạng chuẩn"),
            ("3", "password_hash", "VARCHAR(255)", "-", "NOT NULL", "Mật khẩu đã mã hóa BCrypt (độ an toàn 12)"),
            ("4", "full_name", "VARCHAR(255)", "-", "NOT NULL", "Họ và tên người dùng"),
            ("5", "phone", "VARCHAR(20)", "-", "Nullable", "Số điện thoại liên lạc"),
            ("6", "role", "VARCHAR(20)", "-", "NOT NULL", "Vai trò hệ thống: ADMIN, DRIVER"),
            ("7", "balance", "DECIMAL(12,2)", "-", "NOT NULL", "Số dư ví điện tử của tài xế (mặc định = 0)"),
            ("8", "is_active", "BOOLEAN", "-", "NOT NULL", "Trạng thái hoạt động tài khoản"),
            ("9", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm tạo tài khoản"),
            ("10", "updated_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm cập nhật tài khoản gần nhất")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 2: Vehicles
    p_tbl2 = doc.add_paragraph()
    p_tbl2.add_run("Bảng 2.2: Cấu trúc thực thể vehicles (Phương tiện xe)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "user_id", "BIGINT", "FK", "NOT NULL", "Liên kết với users.id (ON DELETE CASCADE)"),
            ("3", "license_plate", "VARCHAR(20)", "-", "NOT NULL", "Biển số xe ở dạng text thô công khai"),
            ("4", "plate_encrypted", "VARCHAR(255)", "-", "Nullable", "Biển số xe mã hóa bảo mật AES-256"),
            ("5", "vehicle_type", "VARCHAR(50)", "-", "NOT NULL", "Loại xe (ví dụ: CAR, MOTORBIKE)"),
            ("6", "is_default", "BOOLEAN", "-", "NOT NULL", "Đánh dấu phương tiện mặc định"),
            ("7", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm đăng ký phương tiện")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 3: Parking Slots
    p_tbl3 = doc.add_paragraph()
    p_tbl3.add_run("Bảng 2.3: Cấu trúc thực thể parking_slots (Vị trí ô đỗ xe)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "slot_code", "VARCHAR(10)", "Unique", "NOT NULL", "Mã ô đỗ xe (ví dụ: A01, B10)"),
            ("3", "zone", "VARCHAR(10)", "-", "NOT NULL", "Phân khu đỗ xe (Zone A, Zone B)"),
            ("4", "status", "VARCHAR(20)", "-", "NOT NULL", "Trạng thái: AVAILABLE, OCCUPIED, RESERVED..."),
            ("5", "sensor_id", "VARCHAR(100)", "-", "Nullable", "Mã định danh cảm biến kết nối"),
            ("6", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm khởi tạo ô đỗ")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 4: Bookings
    p_tbl4 = doc.add_paragraph()
    p_tbl4.add_run("Bảng 2.4: Cấu trúc thực thể bookings (Lịch sử đặt chỗ)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "user_id", "BIGINT", "FK", "NOT NULL", "Liên kết với users.id"),
            ("3", "vehicle_id", "BIGINT", "FK", "NOT NULL", "Liên kết với vehicles.id"),
            ("4", "slot_id", "BIGINT", "FK", "NOT NULL", "Liên kết với parking_slots.id"),
            ("5", "booking_code", "VARCHAR(50)", "Unique", "NOT NULL", "Mã đặt chỗ ngẫu nhiên"),
            ("6", "qr_code_data", "TEXT", "-", "Nullable", "Dữ liệu JSON mã QR đã ký"),
            ("7", "status", "VARCHAR(20)", "-", "NOT NULL", "Trạng thái: PENDING, CONFIRMED, CHECKED_IN..."),
            ("8", "booked_from", "TIMESTAMP", "-", "NOT NULL", "Thời gian bắt đầu đỗ xe dự kiến"),
            ("9", "booked_until", "TIMESTAMP", "-", "NOT NULL", "Thời gian kết thúc đỗ xe dự kiến"),
            ("10", "checked_in_at", "TIMESTAMP", "-", "Nullable", "Thời điểm xe thực tế vào bãi"),
            ("11", "checked_out_at", "TIMESTAMP", "-", "Nullable", "Thời điểm xe thực tế ra bãi"),
            ("12", "total_amount", "DECIMAL(12,2)", "-", "Nullable", "Tổng chi phí thanh toán cho lượt đỗ"),
            ("13", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm giao dịch đặt chỗ được tạo")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 5: Transactions
    p_tbl5 = doc.add_paragraph()
    p_tbl5.add_run("Bảng 2.5: Cấu trúc thực thể transactions (Giao dịch tài chính)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "booking_id", "BIGINT", "FK", "Nullable", "Liên kết với bookings.id"),
            ("3", "user_id", "BIGINT", "FK", "NOT NULL", "Liên kết với users.id"),
            ("4", "amount", "DECIMAL(12,2)", "-", "NOT NULL", "Số tiền giao dịch (cộng/trừ)"),
            ("5", "payment_method", "VARCHAR(30)", "-", "NOT NULL", "Phương thức thanh toán (mặc định WALLET)"),
            ("6", "payment_status", "VARCHAR(20)", "-", "NOT NULL", "Trạng thái: PENDING, COMPLETED, FAILED"),
            ("7", "transaction_ref", "VARCHAR(100)", "-", "Nullable", "Mã tham chiếu giao dịch ngoài"),
            ("8", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm phát sinh giao dịch")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 6: Device Registry
    p_tbl6 = doc.add_paragraph()
    p_tbl6.add_run("Bảng 2.6: Cấu trúc thực thể device_registry (Đăng ký thiết bị)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "device_uid", "VARCHAR(100)", "Unique", "NOT NULL", "Mã định danh duy nhất của thiết bị"),
            ("3", "device_type", "VARCHAR(30)", "-", "NOT NULL", "Loại thiết bị: ESP32_SENSOR, RPI_GATE"),
            ("4", "location", "VARCHAR(255)", "-", "Nullable", "Vị trí vật lý lắp đặt"),
            ("5", "certificate_cn", "VARCHAR(255)", "-", "Nullable", "Common Name của chứng chỉ số mTLS"),
            ("6", "is_online", "BOOLEAN", "-", "NOT NULL", "Trạng thái trực tuyến"),
            ("7", "last_heartbeat", "TIMESTAMP", "-", "Nullable", "Thời điểm heartbeat cuối cùng"),
            ("8", "firmware_version", "VARCHAR(50)", "-", "Nullable", "Phiên bản phần mềm nhúng hiện tại"),
            ("9", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm đăng ký thiết bị")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # Table 7: Security Audit Log
    p_tbl7 = doc.add_paragraph()
    p_tbl7.add_run("Bảng 2.7: Cấu trúc thực thể security_audit_log (Nhật ký kiểm toán)").bold = True
    add_styled_table(
        doc,
        ["STT", "Tên trường", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"],
        [
            ("1", "id", "BIGSERIAL", "PK", "NOT NULL", "Khóa chính tự tăng"),
            ("2", "user_id", "BIGINT", "FK", "Nullable", "Liên kết với users.id của đối tượng thực hiện"),
            ("3", "action", "VARCHAR(100)", "-", "NOT NULL", "Hành động (LOGIN, GATE_CONTROL,...)"),
            ("4", "resource", "VARCHAR(255)", "-", "Nullable", "Tên thực thể/tài nguyên bị tác động"),
            ("5", "ip_address", "VARCHAR(50)", "-", "Nullable", "Địa chỉ IP nguồn của request gửi lên"),
            ("6", "details", "TEXT", "-", "Nullable", "Mô tả chi tiết dạng chuỗi JSON"),
            ("7", "created_at", "TIMESTAMP", "-", "NOT NULL", "Thời điểm hệ thống ghi nhận sự kiện")
        ],
        [0.5, 1.2, 1.2, 0.6, 1.0, 2.5]
    )

    # 2.2
    h2_2 = doc.add_heading(level=2)
    r2_2 = h2_2.add_run("2.2. Xây dựng tầng REST APIs nghiệp vụ (Spring Boot)")
    r2_2.font.size = Pt(13)
    r2_2.bold = True
    r2_2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Spring Boot 3.2.x được chọn làm nền tảng phát triển Backend Core để tận dụng cơ chế quản lý giao dịch mạnh mẽ và tính bảo mật của Spring Security. "
        "Các API được thiết kế chuẩn RESTful, sử dụng đối tượng ApiResponse<T> đồng nhất để trả dữ liệu về phía client. Dưới đây là bảng danh sách các APIs chi tiết của dự án:"
    )

    # REST API Table
    p_tbl_api = doc.add_paragraph()
    p_tbl_api.add_run("Bảng 2.8: Danh sách các APIs RESTful của hệ thống").bold = True
    add_styled_table(
        doc,
        ["Method", "Endpoint", "Phân quyền", "Mô tả chức năng nghiệp vụ"],
        [
            ("POST", "/api/v1/auth/register", "Public", "Đăng ký tài khoản tài xế mới"),
            ("POST", "/api/v1/auth/login", "Public", "Đăng nhập nhận Access & Refresh Token"),
            ("POST", "/api/v1/auth/refresh", "Public", "Yêu cầu làm mới Access Token"),
            ("POST", "/api/v1/auth/logout", "Authenticated", "Đăng xuất tài khoản, thu hồi token"),
            ("GET", "/api/v1/slots", "ADMIN, DRIVER", "Lấy toàn bộ danh sách ô đỗ xe"),
            ("GET", "/api/v1/slots/available", "ADMIN, DRIVER", "Lấy danh sách các ô đỗ xe còn trống"),
            ("GET", "/api/v1/slots/{id}", "ADMIN, DRIVER", "Lấy thông tin chi tiết của một ô đỗ"),
            ("POST", "/api/v1/bookings", "DRIVER", "Tạo đặt chỗ đỗ xe mới (Pessimistic Lock)"),
            ("GET", "/api/v1/bookings/my", "DRIVER", "Lấy danh sách đặt chỗ của tài xế hiện tại"),
            ("GET", "/api/v1/bookings/{id}", "ADMIN, DRIVER", "Xem thông tin chi tiết một booking"),
            ("POST", "/api/v1/bookings/{id}/cancel", "DRIVER", "Hủy lượt đặt chỗ đỗ xe"),
            ("POST", "/api/v1/bookings/{id}/check-in", "ADMIN", "Xác nhận xe vào bãi"),
            ("POST", "/api/v1/bookings/{id}/check-out", "ADMIN", "Xác nhận xe ra bãi, khấu trừ tiền ví"),
            ("GET", "/api/v1/wallet/balance", "DRIVER", "Lấy số dư tài khoản ví điện tử"),
            ("POST", "/api/v1/wallet/topup", "DRIVER", "Nạp tiền vào tài khoản ví của tài xế"),
            ("GET", "/api/v1/transactions/my", "DRIVER", "Xem lịch sử nạp/trừ tiền"),
            ("GET", "/api/v1/vehicles/my", "DRIVER", "Xem danh sách xe đã đăng ký"),
            ("POST", "/api/v1/vehicles", "DRIVER", "Đăng ký thêm phương tiện mới"),
            ("DELETE", "/api/v1/vehicles/{id}", "DRIVER", "Xóa phương tiện đăng ký"),
            ("GET", "/api/v1/admin/dashboard", "ADMIN", "Lấy thông số thống kê tổng quan bãi xe"),
            ("GET", "/api/v1/admin/revenue", "ADMIN", "Lấy báo cáo doanh thu bãi đỗ theo ngày"),
            ("GET", "/api/v1/admin/users", "ADMIN", "Xem danh sách toàn bộ người dùng"),
            ("PUT", "/api/v1/admin/users/{id}/toggle-active", "ADMIN", "Khóa/mở khóa tài khoản người dùng"),
            ("GET", "/api/v1/admin/devices", "ADMIN", "Giám sát trạng thái thiết bị IoT"),
            ("POST", "/api/v1/admin/gate/{id}/control", "ADMIN", "Điều khiển đóng/mở barrier qua MQTT"),
            ("GET", "/api/v1/admin/audit-logs", "ADMIN", "Truy vấn nhật ký kiểm toán bảo mật"),
            ("POST", "/api/v1/devices/slot-status", "Public / mTLS", "Telemetry cập nhật trạng thái ô đỗ"),
            ("POST", "/api/v1/devices/verify-plate", "Public / mTLS", "RPi gửi biển số xe ANPR xác thực cổng"),
            ("POST", "/api/v1/devices/verify-qr", "Public / mTLS", "RPi gửi dữ liệu QR code xác thực cổng"),
            ("POST", "/api/v1/devices/heartbeat", "Public / mTLS", "Thiết bị gửi tín hiệu duy trì trạng thái")
        ],
        [0.8, 2.2, 1.2, 2.8]
    )

    doc.add_paragraph(
        "Một trong những cơ chế nghiệp vụ quan trọng nhất là tạo đặt chỗ đỗ xe (Create Booking) có áp dụng chống tranh chấp chỗ đỗ (Race Condition). "
        "Khi nhiều tài xế cùng gửi yêu cầu đặt chỗ cho một vị trí đỗ (Parking Slot) duy nhất tại cùng một thời điểm, hệ thống sử dụng khóa bi quan (Pessimistic Lock) mức cơ sở dữ liệu. "
        "Mã nguồn triển khai khóa bi quan tại tầng Repository của Spring Data JPA được thực hiện như sau:"
    )

    # Code block: Pessimistic Lock
    add_code_block(doc, 
        "@Repository\n"
        "public interface ParkingSlotRepository extends JpaRepository<ParkingSlot, Long> {\n"
        "    // Khóa bi quan hàng bản ghi trong Database khi tiến hành giao dịch đặt chỗ\n"
        "    @Lock(LockModeType.PESSIMISTIC_WRITE)\n"
        "    @Query(\"SELECT s FROM ParkingSlot s WHERE s.id = :id\")\n"
        "    Optional<ParkingSlot> findByIdForUpdate(@Param(\"id\") Long id);\n"
        "}"
    )

    doc.add_paragraph(
        "Cơ chế này khóa bản ghi ô đỗ xe ngay khi bắt đầu tiến trình giao dịch đặt chỗ, ngăn không cho các luồng xử lý khác đọc/ghi đè trạng thái của ô đỗ đó cho đến khi transaction kết thúc (commit hoặc rollback). "
        "Nhờ vậy, hiện trạng một ô đỗ bị đặt trùng (Double Booking) hoàn toàn được giải quyết triệt để."
    )

    # 2.3
    h2_3 = doc.add_heading(level=2)
    r2_3 = h2_3.add_run("2.3. Tích hợp kết nối WebSocket STOMP real-time")
    r2_3.font.size = Pt(13)
    r2_3.bold = True
    r2_3.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Để đảm bảo Web Admin Dashboard và Mobile App có thể hiển thị trạng thái các vị trí đỗ xe một cách tức thời ngay khi cảm biến IoT thay đổi trạng thái, hệ thống triển khai giao thức WebSocket STOMP. "
        "Handshake thiết lập kết nối được bảo vệ bởi lớp kiểm tra và validate chữ ký JWT token tại header. Cấu hình WebSocket Broker được cài đặt trong lớp WebSocketConfig như sau:"
    )

    # Code block: WebSocketConfig
    add_code_block(doc,
        "@Configuration\n"
        "@EnableWebSocketMessageBroker\n"
        "public class WebSocketConfig implements WebSocketMessageBrokerConfigurer {\n"
        "    @Override\n"
        "    public void registerStompEndpoints(StompEndpointRegistry registry) {\n"
        "        registry.addEndpoint(\"/ws\").setAllowedOriginPatterns(\"*\").withSockJS();\n"
        "    }\n"
        "    @Override\n"
        "    public void configureMessageBroker(MessageBrokerRegistry registry) {\n"
        "        registry.enableSimpleBroker(\"/topic\");\n"
        "        registry.setApplicationDestinationPrefixes(\"/app\");\n"
        "    }\n"
        "}"
    )

    doc.add_paragraph(
        "Phía Client ReactJS sử dụng thư viện @stomp/stompjs và sockjs-client để bắt tay kết nối và subscribe các sự kiện bãi xe thông qua mã nguồn được tổ chức trong tệp tin websocket.js:"
    )

    # Code block: websocket.js
    add_code_block(doc,
        "import { Client } from '@stomp/stompjs';\n"
        "import SockJS from 'sockjs-client';\n"
        "\n"
        "export const connectWebSocket = (onConnected, onError) => {\n"
        "  const token = localStorage.getItem('token');\n"
        "  stompClient = new Client({\n"
        "    webSocketFactory: () => new SockJS(WS_URL),\n"
        "    connectHeaders: token ? { Authorization: `Bearer ${token}` } : {},\n"
        "    reconnectDelay: 5000,\n"
        "    onConnect: () => {\n"
        "      // Đăng ký lại toàn bộ các topic khi kết nối thành công\n"
        "      subscriptions.forEach(({ topic, callback }) => {\n"
        "        stompClient.subscribe(topic, (msg) => callback(JSON.parse(msg.body)));\n"
        "      });\n"
        "      onConnected?.();\n"
        "    }\n"
        "  });\n"
        "  stompClient.activate();\n"
        "};"
    )

    # 2.4
    h2_4 = doc.add_heading(level=2)
    r2_4 = h2_4.add_run("2.4. Hướng dẫn cài đặt và kết nối Cơ sở dữ liệu (PostgreSQL)")
    r2_4.font.size = Pt(13)
    r2_4.bold = True
    r2_4.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Để phục vụ công tác kiểm thử, phát triển và vận hành hệ thống một cách nhanh chóng, Cơ sở dữ liệu PostgreSQL 16 "
        "được cấu hình chạy trực tiếp thông qua Docker Compose. Điều này đảm bảo tính đồng bộ của môi trường phát triển "
        "và đơn giản hóa các thao tác thiết lập."
    )

    doc.add_paragraph(
        "Các bước cụ thể để khởi chạy và kết nối cơ sở dữ liệu như sau:"
    )

    p_step1 = doc.add_paragraph()
    p_step1.add_run("Bước 1: Khởi chạy PostgreSQL trong Docker").bold = True
    doc.add_paragraph(
        "Mở terminal tại thư mục gốc của dự án (nơi chứa tệp tin docker-compose.yml) và thực hiện lệnh:"
    )
    add_code_block(doc, "docker compose up -d postgres")
    doc.add_paragraph(
        "Lưu ý về cấu hình mạng bảo mật: Nhằm tăng tính an toàn và ngăn chặn các kết nối bên ngoài trực tiếp tấn công database, "
        "mạng backend-net được cài đặt cấu hình internal: true. Để có thể truy cập và kiểm tra dữ liệu từ máy host thông qua các ứng dụng "
        "GUI (như DataGrip), dịch vụ postgres được cấu hình kết nối đồng thời vào mạng smartparking-frontend (frontend-net). Nhờ đó, "
        "cổng 5432 của PostgreSQL được ánh xạ ra localhost thành công."
    )

    p_step2 = doc.add_paragraph()
    p_step2.add_run("Bước 2: Cấu hình thông tin kết nối").bold = True
    doc.add_paragraph(
        "Các thông số môi trường của cơ sở dữ liệu được quản lý tập trung tại tệp tin .env:"
    )
    add_code_block(doc,
        "POSTGRES_DB=smartparking\n"
        "POSTGRES_USER=smartparking\n"
        "POSTGRES_PASSWORD=smartparking\n"
        "POSTGRES_PORT=5432"
    )

    p_step3 = doc.add_paragraph()
    p_step3.add_run("Bước 3: Hướng dẫn kết nối bằng JetBrains DataGrip").bold = True
    doc.add_paragraph(
        "Nhà phát triển có thể dễ dàng quản lý dữ liệu trực quan bằng các bước sau trên DataGrip:\n"
        "1. Trên thanh công cụ Database (phía bên trái màn hình), nhấn biểu tượng '+' -> Data Source -> PostgreSQL.\n"
        "2. Thiết lập các thông số kết nối tương ứng:\n"
        "   - Host: localhost (hoặc 127.0.0.1)\n"
        "   - Port: 5432\n"
        "   - User: smartparking\n"
        "   - Password: smartparking\n"
        "   - Database: smartparking\n"
        "3. Tải Driver kết nối: Nhấn vào liên kết 'Download' ở cạnh dòng chữ cảnh báo thiếu driver phía dưới cửa sổ cài đặt.\n"
        "4. Kiểm tra kết nối: Nhấn vào nút 'Test Connection'. Khi hệ thống thông báo kết nối thành công (màu xanh lá), nhấn 'OK' hoặc 'Apply' để hoàn tất kết nối.\n"
        "5. Sau khi kết nối, cấu trúc cơ sở dữ liệu (schemas, tables) và dữ liệu mẫu (seed data) được khởi tạo tự động từ script init.sql sẽ hiển thị sẵn sàng để truy vấn."
    )

    doc.add_page_break()

    # ==================== CHƯƠNG 3 ====================
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("CHƯƠNG 3: GIẢI PHÁP BẢO MẬT KIẾN TRÚC HỆ THỐNG")
    r3.font.size = Pt(15)
    r3.bold = True
    r3.font.color.rgb = RGBColor(0, 0, 0)

    # 3.1
    h3_1 = doc.add_heading(level=2)
    r3_1 = h3_1.add_run("3.1. Xác thực phân quyền không trạng thái (JWT Stateless Auth & Blacklist)")
    r3_1.font.size = Pt(13)
    r3_1.bold = True
    r3_1.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Tầng bảo mật xác thực người dùng được thiết kế dựa trên cơ chế JWT Stateless. "
        "Mỗi khi đăng nhập thành công, hệ thống sinh ra cặp Access Token (hết hạn sau 15 phút, chứa claims: userId, email, role) và Refresh Token (hết hạn sau 7 ngày) được ký bằng thuật toán HMAC-SHA256."
    )
    doc.add_paragraph(
        "Để hỗ trợ tính năng đăng xuất an toàn (Secure Logout) trong kiến trúc stateless, hệ thống sử dụng bộ nhớ cache Redis làm kho lưu trữ Token Blacklist. "
        "Khi người dùng gọi API đăng xuất, token hiện tại được nạp vào Redis với thời gian sống (TTL) tương ứng với thời hạn hết hạn còn lại của chính token đó. "
        "Bộ lọc JwtAuthenticationFilter chặn mọi request để kiểm tra tính hợp lệ của token và đối soát với blacklist của Redis:"
    )

    # Code block: JwtAuthenticationFilter
    add_code_block(doc,
        "// Trích đoạn mã nguồn xử lý lọc token và đối chiếu Blacklist Redis\n"
        "@Component\n"
        "public class JwtAuthenticationFilter extends OncePerRequestFilter {\n"
        "    private final JwtService jwtService;\n"
        "    private final StringRedisTemplate redisTemplate;\n"
        "\n"
        "    @Override\n"
        "    protected void doFilterInternal(HttpServletRequest request, HttpServletResponse response, \n"
        "                                    FilterChain filterChain) throws ServletException, IOException {\n"
        "        final String authHeader = request.getHeader(\"Authorization\");\n"
        "        if (authHeader == null || !authHeader.startsWith(\"Bearer \")) {\n"
        "            filterChain.doFilter(request, response);\n"
        "            return;\n"
        "        }\n"
        "        String jwt = authHeader.substring(7);\n"
        "        // Kiểm tra xem token đã bị đăng xuất chưa (Blacklist Redis)\n"
        "        Boolean isBlacklisted = redisTemplate.hasKey(\"blacklist:\" + jwt);\n"
        "        if (Boolean.TRUE.equals(isBlacklisted)) {\n"
        "            response.setStatus(HttpServletResponse.SC_UNAUTHORIZED);\n"
        "            response.setContentType(\"application/json\");\n"
        "            response.getWriter().write(\"{\\\"success\\\":false,\\\"message\\\":\\\"Token invalid (logged out)\\\"}\");\n"
        "            return;\n"
        "        }\n"
        "        // Xử lý xác thực người dùng nếu token hợp lệ...\n"
        "        filterChain.doFilter(request, response);\n"
        "    }\n"
        "}"
    )

    # 3.2
    h3_2 = doc.add_heading(level=2)
    r3_2 = h3_2.add_run("3.2. Mã hóa dữ liệu lưu trữ (AES-256 License Plate Encryption)")
    r3_2.font.size = Pt(13)
    r3_2.bold = True
    r3_2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Biển số xe là thông tin định danh cá nhân nhạy cảm, cần được bảo vệ trước các nguy cơ tấn công dò quét hoặc rò rỉ dữ liệu thô từ database. "
        "Hệ thống triển khai dịch vụ CryptoService để mã hóa biển số xe bằng thuật toán mã hóa đối xứng mạnh AES-256 ở chế độ CBC (Cipher Block Chaining) và cơ chế padding PKCS5Padding."
    )
    doc.add_paragraph(
        "Mỗi lần mã hóa, hệ thống sinh ngẫu nhiên một Vector khởi tạo (Initialization Vector - IV) dài 16 bytes nhằm đảm bảo tính ngẫu nhiên. "
        "IV được ghép trực tiếp vào đầu chuỗi ciphertext trước khi lưu trữ vào cột plate_encrypted trong Database. "
        "Đoạn mã mã hóa và giải mã trong CryptoService được xây dựng như sau:"
    )

    # Code block: CryptoService
    add_code_block(doc,
        "public class CryptoService {\n"
        "    private static final String ALGORITHM = \"AES/CBC/PKCS5Padding\";\n"
        "    private static final int IV_SIZE = 16;\n"
        "    private final SecureRandom secureRandom = new SecureRandom();\n"
        "\n"
        "    public String encrypt(String plainText) {\n"
        "        try {\n"
        "            byte[] iv = new byte[IV_SIZE];\n"
        "            secureRandom.nextBytes(iv); // Sinh ngẫu nhiên IV\n"
        "            IvParameterSpec ivSpec = new IvParameterSpec(iv);\n"
        "            Cipher cipher = Cipher.getInstance(ALGORITHM);\n"
        "            cipher.init(Cipher.ENCRYPT_MODE, getSecretKeySpec(), ivSpec);\n"
        "            byte[] cipherTextBytes = cipher.doFinal(plainText.getBytes(StandardCharsets.UTF_8));\n"
        "            // Ghép IV vào trước bản mã hóa\n"
        "            byte[] combined = new byte[IV_SIZE + cipherTextBytes.length];\n"
        "            System.arraycopy(iv, 0, combined, 0, IV_SIZE);\n"
        "            System.arraycopy(cipherTextBytes, 0, combined, IV_SIZE, cipherTextBytes.length);\n"
        "            return Base64.getEncoder().encodeToString(combined);\n"
        "        } catch (Exception e) { throw new RuntimeException(\"Encryption error\", e); }\n"
        "    }\n"
        "    // Triển khai giải mã tương tự bằng cách tách 16 byte đầu làm IV...\n"
        "}"
    )

    # 3.3
    h3_3 = doc.add_heading(level=2)
    r3_3 = h3_3.add_run("3.3. Cơ chế chống tấn công phát lại (Redis Nonce Anti-Replay)")
    r3_3.font.size = Pt(13)
    r3_3.bold = True
    r3_3.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Tấn công phát lại (Replay Attack) xảy ra khi kẻ tấn công bắt gói tin điều khiển barie cổng hợp lệ được truyền đi trước đó và gửi lại lệnh đó cho thiết bị để mở barrier bất hợp pháp. "
        "Để chống lại hình thức tấn công này, hệ thống triển khai cơ chế kết hợp Nonce (Number used once) và Timestamp trên bộ nhớ đệm Redis:"
    )
    doc.add_paragraph(
        "Khi Admin gửi lệnh mở barie qua cổng quản trị, hệ thống sinh ra một mã UUID làm Nonce, kèm theo mốc thời gian Timestamp hiện tại của server, ký mã hóa gói lệnh điều khiển và truyền đi qua MQTT. "
        "Thiết bị nhận lệnh (Raspberry Pi điều khiển cổng) hoặc Backend xử lý sự kiện sẽ kiểm tra chênh lệch thời gian tối đa 5 giây và trạng thái Nonce trong Redis:"
    )

    # Code block: NonceServiceImpl
    add_code_block(doc,
        "@Service\n"
        "public class NonceServiceImpl implements NonceService {\n"
        "    @Autowired\n"
        "    private StringRedisTemplate redisTemplate;\n"
        "    @Value(\"${app.nonce-max-age-seconds:5}\")\n"
        "    private long maxAgeSeconds;\n"
        "\n"
        "    @Override\n"
        "    public void validateNonce(String nonce, long timestamp) {\n"
        "        long currentMillis = System.currentTimeMillis();\n"
        "        long diffMillis = Math.abs(currentMillis - timestamp);\n"
        "        // 1. Kiểm tra giới hạn lệch thời gian\n"
        "        if (diffMillis > maxAgeSeconds * 1000) {\n"
        "            throw new ReplayAttackException(\"Request timestamp is outside the allowed window\");\n"
        "        }\n"
        "        // 2. Sử dụng lệnh SETNX nguyên tử trên Redis để lưu nonce\n"
        "        String redisKey = \"nonce:\" + nonce;\n"
        "        Boolean isAbsent = redisTemplate.opsForValue().setIfAbsent(redisKey, \"1\", \n"
        "                                               maxAgeSeconds * 2, TimeUnit.SECONDS);\n"
        "        if (Boolean.FALSE.equals(isAbsent)) {\n"
        "            throw new ReplayAttackException(\"Replay attack detected: Nonce already used\");\n"
        "        }\n"
        "    }\n"
        "}"
    )

    # 3.4
    h3_4 = doc.add_heading(level=2)
    r3_4 = h3_4.add_run("3.4. Kiểm toán và ghi nhật ký bảo mật (Security Audit Logging & Spring AOP)")
    r3_4.font.size = Pt(13)
    r3_4.bold = True
    r3_4.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Mọi hoạt động nghiệp vụ quan trọng hay các hành vi tác động vào cấu hình bảo mật đều được ghi nhận tự động vào bảng security_audit_log. "
        "Hệ thống áp dụng lập trình hướng khía cạnh (Aspect Oriented Programming - AOP) để tách biệt logic ghi nhật ký kiểm toán khỏi logic xử lý nghiệp vụ chính. "
        "Lớp AuditAspect được triển khai để tự động bắt và lưu trữ các hành động bảo mật cốt lõi:"
    )

    # Code block: AuditAspect
    add_code_block(doc,
        "@Aspect\n"
        "@Component\n"
        "@RequiredArgsConstructor\n"
        "public class AuditAspect {\n"
        "    private final AuditService auditService;\n"
        "    private final UserRepository userRepository;\n"
        "\n"
        "    // Ghi nhật ký khi người dùng đăng nhập thành công\n"
        "    @AfterReturning(pointcut = \"execution(* com.smartparking.controller.AuthController.login(..)) && args(loginRequest)\", returning = \"result\")\n"
        "    public void logLoginSuccess(LoginRequest loginRequest, Object result) {\n"
        "        userRepository.findByEmail(loginRequest.getEmail()).ifPresent(user -> {\n"
        "            auditService.log(user.getId(), \"LOGIN_SUCCESS\", \"users\", getClientIp(), \"User logged in successfully\");\n"
        "        });\n"
        "    }\n"
        "    // Ghi nhật ký khi đăng nhập thất bại\n"
        "    @AfterThrowing(pointcut = \"execution(* com.smartparking.controller.AuthController.login(..)) && args(loginRequest)\", throwing = \"ex\")\n"
        "    public void logLoginFailure(LoginRequest loginRequest, Throwable ex) {\n"
        "        Long userId = userRepository.findByEmail(loginRequest.getEmail()).map(User::getId).orElse(null);\n"
        "        auditService.log(userId, \"LOGIN_FAILED\", \"users\", getClientIp(), \"Failed attempt: \" + ex.getMessage());\n"
        "    }\n"
        "    // Ghi nhật ký khi Admin can thiệp điều khiển barrier\n"
        "    @AfterReturning(\"execution(* com.smartparking.controller.AdminController.controlGate(..)) && args(gateId, request)\")\n"
        "    public void logGateControl(String gateId, GateControlRequest request) {\n"
        "        auditService.log(getCurrentUserId(), \"GATE_CONTROL\", \"devices\", getClientIp(), \"Controlled gate: \" + gateId);\n"
        "    }\n"
        "    // Ghi nhận IP thực qua Header X-Forwarded-For của Nginx...\n"
        "}"
    )

    # 3.5
    h3_5 = doc.add_heading(level=2)
    r3_5 = h3_5.add_run("3.5. Giới hạn tần suất yêu cầu (Rate Limiting với Bucket4j)")
    r3_5.font.size = Pt(13)
    r3_5.bold = True
    r3_5.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Nhằm ngăn chặn các cuộc tấn công Brute-Force hoặc dò tìm thông tin đăng nhập tài khoản diện rộng, hệ thống triển khai bộ giới hạn tần suất (Rate Limiting) tại endpoint đăng nhập. "
        "Giải thuật Token Bucket được chọn làm mô hình hoạt động chính và hiện thực hóa thông qua thư viện Bucket4j của Java."
    )
    doc.add_paragraph(
        "Mỗi địa chỉ email khi thực hiện gửi request đăng nhập sẽ được cấp phát một bộ bucket có kích thước tối đa là 10 tokens. "
        "Mỗi lượt đăng nhập tiêu thụ 1 token. Bộ bucket áp dụng cơ chế Refill.greedy hồi phục toàn bộ 10 tokens sau mỗi khoảng thời gian 5 phút. "
        "Lớp RateLimitAspect thực hiện chặn đầu yêu cầu đăng nhập như sau:"
    )

    # Code block: RateLimitAspect
    add_code_block(doc,
        "@Aspect\n"
        "@Component\n"
        "public class RateLimitAspect {\n"
        "    private final Map<String, Bucket> cache = new ConcurrentHashMap<>();\n"
        "\n"
        "    @Before(\"execution(* com.smartparking.controller.AuthController.login(..)) && args(request)\")\n"
        "    public void rateLimitLogin(LoginRequest request) {\n"
        "        String email = request.getEmail();\n"
        "        Bucket bucket = cache.computeIfAbsent(email, k -> createNewBucket());\n"
        "        if (!bucket.tryConsume(1)) {\n"
        "            long nanosToWait = bucket.estimateAbilityToConsume(1).getNanosToWait();\n"
        "            long secondsToWait = Math.max(1, TimeUnit.NANOSECONDS.toSeconds(nanosToWait));\n"
        "            throw new TooManyRequestsException(\"Đăng nhập quá nhanh. Vui lòng thử lại sau.\", secondsToWait);\n"
        "        }\n"
        "    }\n"
        "    private Bucket createNewBucket() {\n"
        "        return Bucket.builder()\n"
        "                .addLimit(Bandwidth.classic(10, Refill.greedy(10, Duration.ofMinutes(5))))\n"
        "                .build();\n"
        "    }\n"
        "}"
    )

    # 3.6
    h3_6 = doc.add_heading(level=2)
    r3_6 = h3_6.add_run("3.6. Cấu hình cổng bảo mật API Gateway và SSL Termination (Nginx)")
    r3_6.font.size = Pt(13)
    r3_6.bold = True
    r3_6.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Nginx được cấu hình như một reverse proxy và API Gateway duy nhất giao tiếp với bên ngoài. "
        "Mọi truy cập HTTP đều được cấu hình tự động chuyển hướng (HTTP 301 Redirect) sang cổng HTTPS 443. "
        "Tại cổng 443, Nginx đóng vai trò SSL Termination sử dụng giao thức mật mã hiện đại TLS 1.2 và TLS 1.3 với các bộ ciphers bảo mật mạnh. "
        "Nginx cũng trực tiếp cấu hình rate limiting ở mức IP client và chặn truy cập trái phép vào các endpoint nhạy cảm (như Actuator). "
        "Dưới đây là một phần cấu hình máy chủ SSL trong nginx.conf:"
    )

    # Code block: nginx.conf server block
    add_code_block(doc,
        "server {\n"
        "    listen 443 ssl http2;\n"
        "    ssl_certificate     /etc/nginx/certs/server.crt;\n"
        "    ssl_certificate_key /etc/nginx/certs/server.key;\n"
        "    ssl_protocols TLSv1.2 TLSv1.3;\n"
        "    \n"
        "    # Thêm tiêu đề an toàn phòng chống XSS, Clickjacking, và HSTS\n"
        "    add_header X-Frame-Options \"DENY\" always;\n"
        "    add_header X-Content-Type-Options \"nosniff\" always;\n"
        "    add_header Strict-Transport-Security \"max-age=31536000; includeSubDomains\" always;\n"
        "\n"
        "    # Định tuyến Proxy API vào Spring Boot Backend\n"
        "    location /api/ {\n"
        "        limit_req zone=api burst=20 nodelay; # Giới hạn request mức Nginx\n"
        "        proxy_pass http://backend;\n"
        "        proxy_set_header X-Real-IP $remote_addr;\n"
        "        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n"
        "    }\n"
        "}"
    )

    doc.add_page_break()

    # ==================== CHƯƠNG 4 ====================
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("CHƯƠNG 4: LẬP TRÌNH GIAO DIỆN CLIENT & THIẾT BỊ IoT")
    r4.font.size = Pt(15)
    r4.bold = True
    r4.font.color.rgb = RGBColor(0, 0, 0)

    # 4.1
    h4_1 = doc.add_heading(level=2)
    r4_1 = h4_1.add_run("4.1. Cấu trúc tổ chức mã nguồn Monorepo")
    r4_1.font.size = Pt(13)
    r4_1.bold = True
    r4_1.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Hệ thống Smart Parking được tổ chức dưới dạng cấu trúc Monorepo thống nhất tại thư mục gốc, giúp quản lý toàn bộ các thành phần của hệ thống (Backend, Web Dashboard, Mobile App, IoT Hardware Code, Docker Infrastructure) một cách tập trung. "
        "Cấu trúc thư mục mã nguồn được phân bổ như sau:"
    )

    # Code block: Monorepo Structure
    add_code_block(doc,
        "Smart Parking System/\n"
        "├── backend/                  # Mã nguồn Spring Boot Java Backend\n"
        "│   ├── src/main/java/com/smartparking/\n"
        "│   └── pom.xml\n"
        "├── web/                      # Mã nguồn Admin Web Dashboard (Vite + ReactJS)\n"
        "│   ├── src/                  # Components, Pages, Stores (Zustand)\n"
        "│   └── package.json\n"
        "├── mobile/                   # Mã nguồn Mobile App cho tài xế (Flutter)\n"
        "│   ├── lib/                  # Screens, Providers, Services\n"
        "│   └── pubspec.yaml\n"
        "├── iot/                      # Mã nguồn phần cứng IoT\n"
        "│   ├── esp32/                # Firmware cảm biến đỗ xe (PlatformIO C++)\n"
        "│   └── raspberry-pi/         # Script Python xử lý ANPR & quét QR tại cổng\n"
        "├── docker/                   # Hạ tầng Docker\n"
        "│   ├── nginx/                # File cấu hình Nginx & SSL Certs\n"
        "│   ├── mosquitto/            # File cấu hình MQTT Broker & ACL & Certs\n"
        "│   ├── certs/                # Scripts tự sinh CA và chứng chỉ số mTLS\n"
        "│   └── postgres/             # Tập lệnh khởi tạo CSDL (init.sql)\n"
        "└── docker-compose.yml        # Docker Orchestration định cấu hình toàn hệ thống"
    )

    # 4.2
    h4_2 = doc.add_heading(level=2)
    r4_2 = h4_2.add_run("4.2. Giao diện & Xử lý Ứng dụng di động tài xế (Flutter Mobile Client)")
    r4_2.font.size = Pt(13)
    r4_2.bold = True
    r4_2.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Ứng dụng Flutter được phát triển dành riêng cho đối tượng tài xế nhằm mục đích tra cứu sơ đồ bãi xe, đặt chỗ và thanh toán. "
        "Giao diện ứng dụng được thiết kế theo các tiêu chuẩn thẩm mỹ hiện đại, sử dụng tông màu Indigo và Teal. Dưới đây là bảng mô tả chi tiết các phân hệ màn hình chính:"
    )

    # Screens Table for Mobile
    p_tbl_mobile = doc.add_paragraph()
    p_tbl_mobile.add_run("Bảng 4.1: Các màn hình giao diện của Driver Mobile App").bold = True
    add_styled_table(
        doc,
        ["Tên màn hình", "Chức năng nghiệp vụ chính", "Các thành phần UI & Tiêu chuẩn thẩm mỹ"],
        [
            ("Login / Register", "Xác thực người dùng", "Form đăng nhập với gradient nền, bo góc mượt, check validation đầy đủ."),
            ("Parking Map (Tab)", "Theo dõi bãi đỗ xe", "Bản đồ lưới ô đỗ xe 2D. Đổi màu real-time nhờ kết nối WebSocket STOMP."),
            ("Bookings (Tab)", "Vé đỗ xe và Lịch sử", "Xem vé đang kích hoạt (có mã QR vé), danh sách vé đã đỗ hoặc đã hủy."),
            ("Wallet (Tab)", "Thanh toán & Số dư", "Thẻ ví điện tử ảo, hiển thị số dư VND, nạp tiền và lịch sử giao dịch ví."),
            ("Profile (Tab)", "Quản lý cá nhân", "Quản lý thông tin tài khoản, đổi mật khẩu và quản lý đăng ký phương tiện xe.")
        ],
        [1.5, 2.5, 3.0]
    )

    doc.add_paragraph(
        "Một trong những kỹ thuật quan trọng nhất trong việc lập trình Client Mobile an toàn là tự động làm mới Access Token (Auto-refresh JWT) khi token hết hạn. "
        "Lớp ApiService trong Flutter chặn mọi phản hồi lỗi HTTP 401 Unauthorized để gửi Refresh Token làm mới Access Token một cách tự động trước khi thử lại request ban đầu của người dùng:"
    )

    # Code block: Auto refresh token Dart
    add_code_block(doc,
        "// Lớp ApiService chặn và tự động refresh token trong Flutter\n"
        "Future<http.Response> get(String path) async {\n"
        "  final url = Uri.parse('${AppConstants.apiBaseUrl}$path');\n"
        "  final headers = await _getHeaders();\n"
        "  var response = await _client.get(url, headers: headers);\n"
        "  \n"
        "  if (response.statusCode == 401) {\n"
        "    // Nếu token hết hạn, tự động gọi hàm refresh\n"
        "    final refreshed = await _attemptTokenRefresh();\n"
        "    if (refreshed) {\n"
        "      // Thử lại request ban đầu với headers chứa token mới\n"
        "      final retryHeaders = await _getHeaders();\n"
        "      response = await _client.get(url, headers: retryHeaders);\n"
        "    }\n"
        "  }\n"
        "  return response;\n"
        "}"
    )

    # 4.3
    h4_3 = doc.add_heading(level=2)
    r4_3 = h4_3.add_run("4.3. Trang điều khiển giám sát trung tâm (ReactJS Admin Dashboard)")
    r4_3.font.size = Pt(13)
    r4_3.bold = True
    r4_3.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Trang quản trị trung tâm được phát triển bằng ReactJS 18 và thư viện giao diện Ant Design 5. "
        "Hệ thống cung cấp các chức năng giám sát chuyên sâu cho Admin bãi xe. Danh sách các trang được tổ chức như sau:"
    )

    # Screens Table for Web
    p_tbl_web = doc.add_paragraph()
    p_tbl_web.add_run("Bảng 4.2: Các phân hệ trang quản trị của Admin Web Dashboard").bold = True
    add_styled_table(
        doc,
        ["Tên trang", "Mô tả chức năng", "Dữ liệu & Kết nối thời gian thực"],
        [
            ("Login Page", "Đăng nhập xác thực tài khoản quản trị", "Lưu token vào LocalStorage, chuyển hướng an toàn qua PrivateRoute."),
            ("Dashboard Page", "Bảng thống kê vận hành bãi xe", "Hiển thị biểu đồ tròn sử dụng (Pie Chart), doanh thu và số lượng ô đỗ trống/bận."),
            ("Parking Map Page", "Sơ đồ bãi xe real-time & Barrier", "Hiển thị màu sắc thực tế ô đỗ thay đổi live qua WebSocket và cho phép admin mở cổng khẩn cấp."),
            ("Revenue Page", "Báo cáo tài chính doanh thu", "Thống kê doanh thu theo khoảng thời gian chọn, vẽ biểu đồ cột doanh thu hàng ngày."),
            ("Users Page", "Quản lý thông tin tài xế", "Danh sách người dùng, hỗ trợ tìm kiếm, phân trang và khóa/mở khóa tài khoản tài xế."),
            ("Devices Page", "Giám sát thiết bị ngoại vi IoT", "Hiển thị UID, IP, loại thiết bị, phiên bản firmware và trạng thái online/offline của ESP32/RPi."),
            ("Audit Logs Page", "Tra cứu nhật ký kiểm toán bảo mật", "Bảng dữ liệu ghi nhận hoạt động bảo mật, hỗ trợ lọc theo hành động và xem chi tiết cấu trúc JSON.")
        ],
        [1.5, 2.5, 3.0]
    )

    doc.add_paragraph(
        "Tất cả các dữ liệu phiên làm việc và trạng thái bãi xe của Client được quản lý tập trung và phân tách logic bằng thư viện Zustand store (như authStore, slotStore), bảo đảm hiệu năng render và cấu trúc code sạch sẽ."
    )

    # 4.4
    h4_4 = doc.add_heading(level=2)
    r4_4 = h4_4.add_run("4.4. Lập trình thiết bị IoT, Barrier & Hạ tầng mTLS Mosquitto")
    r4_4.font.size = Pt(13)
    r4_4.bold = True
    r4_4.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Bảo mật truyền thông thiết bị IoT là ưu tiên hàng đầu của hệ thống. "
        "Mọi kết nối của thiết bị phần cứng tới Mosquitto Broker đều phải qua cổng bảo mật TLS 8883, bắt buộc xác thực hai chiều mTLS (Mutual TLS). "
        "Hệ thống chứng chỉ số được xây dựng tự động hóa hoàn toàn thông qua tệp tin kịch bản bash generate_certs.sh:"
    )

    # Code block: generate_certs.sh steps
    add_code_block(doc,
        "# Quy trình sinh chứng chỉ số tự động bằng OpenSSL trong generate_certs.sh:\n"
        "# 1. Tạo Root CA nội bộ tự ký (valid 10 năm)\n"
        "openssl req -new -x509 -key ca.key -out ca.crt -days 3650 -subj \"/CN=SmartParking-CA\"\n"
        "# 2. Sinh khóa và ký chứng chỉ số cho Mosquitto Broker\n"
        "openssl req -new -key broker.key -out broker.csr -subj \"/CN=mosquitto\"\n"
        "openssl x509 -req -in broker.csr -CA ca.crt -CAkey ca.key -out broker.crt -days 825\n"
        "# 3. Sinh khóa và ký chứng chỉ số client riêng biệt cho từng thiết bị ngoại vi\n"
        "openssl req -new -key esp32_slot_a01.key -out esp32_slot_a01.csr -subj \"/CN=esp32_slot_a01\"\n"
        "openssl x509 -req -in esp32_slot_a01.csr -CA ca.crt -CAkey ca.key -out esp32_slot_a01.crt -days 825"
    )
    
    doc.add_paragraph(
        "Chứng chỉ CA được import vào ESP32 để ESP32 tin tưởng Broker. "
        "Đồng thời, ESP32 gửi chứng chỉ Client của nó cho Broker xác thực. "
        "Để ngăn chặn một thiết bị IoT bị tấn công chiếm quyền có thể can thiệp dữ liệu của thiết bị khác, Mosquitto Broker cấu hình tệp tin ACL (Access Control List) phân quyền nghiêm ngặt dựa trên Common Name (CN) của chứng chỉ số:"
    )

    # Code block: ACL sample
    add_code_block(doc,
        "# Cấu hình ACL Mosquitto phân quyền thiết bị\n"
        "user esp32_slot_a01\n"
        "topic write parking/slots/a01/status\n"
        "topic read parking/slots/a01/command\n"
        "topic write parking/devices/esp32_slot_a01/heartbeat\n"
        "\n"
        "user rpi_gate1\n"
        "topic write parking/gates/gate1/event\n"
        "topic read parking/gates/gate1/control\n"
        "topic write parking/devices/rpi_gate1/heartbeat"
    )

    doc.add_paragraph(
        "Nhờ cấu hình này, thiết bị `esp32_slot_a01` chỉ được phép publish lên topic trạng thái của đúng ô đỗ A01 và subscribe topic chỉ thị đèn của ô đỗ A01. "
        "Mọi hành vi gửi/nhận tin đến các vị trí đỗ khác đều bị Broker từ chối trực tiếp ở lớp giao thức, triệt tiêu nguy cơ lan rộng lỗ hổng bảo mật khi có thiết bị phần cứng bị xâm nhập vật lý."
    )

    doc.add_paragraph(
        "Đối với cổng kiểm soát ra vào bãi đỗ (Gate Controller), một chương trình Python chạy giả lập trên PC/Raspberry Pi chịu trách nhiệm xử lý webcam thu hình ảnh, nhận dạng biển số (YOLOv8 + EasyOCR), quét mã QR (pyzbar), giao tiếp MQTT mTLS và xác thực mã nonce với Backend qua API."
    )
    doc.add_paragraph(
        "Nhằm bảo vệ hệ thống trước tấn công phát lại (Replay attack), khi thiết bị nhận được lệnh mở/đóng cổng qua MQTT topic 'parking/gates/{gateId}/control', nó thực hiện xác thực độ lệch thời gian cục bộ (dưới 5 giây) và gửi yêu cầu POST HTTP kèm theo nonce và timestamp ngược lại Backend để xác minh tính hợp lệ thông qua lớp MQTTHandler:"
    )

    # Code block: RPi MQTT Handler verify-nonce
    add_code_block(doc,
        "# Trích đoạn xử lý xác thực nonce và timestamp cục bộ + remote trong mqtt_handler.py\n"
        "def on_message(self, client, userdata, msg):\n"
        "    try:\n"
        "        payload = json.loads(msg.payload.decode('utf-8'))\n"
        "        action = payload.get(\"action\")\n"
        "        nonce = payload.get(\"nonce\")\n"
        "        timestamp = payload.get(\"timestamp\")\n"
        "\n"
        "        # 1. Kiểm tra độ lệch thời gian cục bộ (5 giây)\n"
        "        current_time_ms = int(time.time() * 1000)\n"
        "        if abs(current_time_ms - timestamp) > 5000:\n"
        "            logger.warning(\"[REPLAY ATTACK] Sai lệch thời gian quá lớn. Bỏ qua lệnh!\")\n"
        "            return\n"
        "\n"
        "        # 2. Gọi API Backend verify-nonce để xác thực và lưu vết chống Replay\n"
        "        verify_url = f\"{API_BASE_URL}/devices/verify-nonce\"\n"
        "        response = requests.post(verify_url, params={\"nonce\": nonce, \"timestamp\": timestamp}, timeout=3.0)\n"
        "        \n"
        "        if response.status_code == 200:\n"
        "            logger.info(\"Xác thực nonce thành công. Tiến hành điều khiển barie.\")\n"
        "            if action == \"OPEN\":\n"
        "                self.gate_controller.open_gate()\n"
        "            elif action == \"CLOSE\":\n"
        "                self.gate_controller.close_gate()\n"
        "        else:\n"
        "            logger.error(\"Backend từ chối xác thực nonce. Lệnh không hợp lệ hoặc đã bị replay!\")\n"
        "    except Exception as e:\n"
        "        logger.error(f\"Lỗi xử lý tin nhắn MQTT: {e}\")"
    )

    doc.add_paragraph(
        "Chương trình chạy chính main.py duy trì luồng đọc hình ảnh từ camera, phân tách xử lý song song giữa nhận dạng biển số ANPR và quét mã QR, đồng thời áp dụng cơ chế thời gian chờ (cooldown 10 giây) sau mỗi lần kích hoạt cổng thành công để tránh lặp lệnh xác thực:"
    )

    # Code block: RPi Main Loop capture
    add_code_block(doc,
        "# Trích đoạn vòng lặp xử lý hình ảnh camera trong main.py\n"
        "while self.running:\n"
        "    ret, frame = cap.read()\n"
        "    if not ret: continue\n"
        "\n"
        "    current_time = time.time()\n"
        "    is_in_cooldown = (current_time - self.last_detection_time) < self.cooldown_period\n"
        "    self.draw_overlay(frame, self.gate_controller.is_open(), is_in_cooldown)\n"
        "\n"
        "    if not is_in_cooldown:\n"
        "        # 1. Quét QR code (ưu tiên hàng đầu, nhẹ CPU)\n"
        "        qr_data = self.qr_scanner.scan_from_frame(frame)\n"
        "        if qr_data:\n"
        "            self.mqtt_handler.publish_event(\"QR_SCANNED\", qr_data)\n"
        "            self.last_detection_time = current_time\n"
        "\n"
        "        # 2. Quét biển số xe ANPR (chạy mỗi 5 frames để tiết kiệm tài nguyên)\n"
        "        elif self.anpr_engine and frame_count % 5 == 0:\n"
        "            cropped, bbox = self.anpr_engine.detect_plate(frame)\n"
        "            if cropped is not None:\n"
        "                plate, conf = self.anpr_engine.recognize_text(cropped)\n"
        "                if plate and len(plate) >= 5:\n"
        "                    self.mqtt_handler.publish_event(\"PLATE_DETECTED\", plate)\n"
        "                    self.last_detection_time = current_time\n"
        "\n"
        "    cv2.imshow(\"Smart Parking - Gate Controller Simulation\", frame)\n"
        "    if cv2.waitKey(1) & 0xFF == ord('q'):\n"
        "        self.running = False"
    )

    doc.add_page_break()

    # ==================== KẾT LUẬN ====================
    h_concl = doc.add_heading(level=1)
    r_concl = h_concl.add_run("KẾT LUẬN & HƯỚNG PHÁT TRIỂN")
    r_concl.font.size = Pt(15)
    r_concl.bold = True
    r_concl.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph(
        "Đề tài đã hoàn thành việc xây dựng toàn bộ kiến trúc hệ thống Smart Parking hoàn chỉnh, giải quyết tốt bài toán nghiệp vụ đỗ xe và đặc biệt là áp dụng thành công các giải pháp bảo mật nhiều lớp từ phần cứng lên phần mềm. "
        "Hệ thống đạt được độ phản hồi dữ liệu thời gian thực cao nhờ WebSocket và MQTT qua mạng kết nối không trạng thái stateless an toàn."
    )
    doc.add_paragraph(
        "Hướng phát triển tiếp theo của đề tài tập trung vào việc:\n"
        "1. Tích hợp AI Camera thông minh tại cổng ra vào để nhận diện khuôn mặt tài xế song song với biển số xe nâng cao bảo mật.\n"
        "2. Áp dụng công nghệ Blockchain trong lưu trữ Transaction History và Security Audit Logs nhằm đảm bảo tính bất biến (immutability) tuyệt đối, chống giả mạo hay xóa nhật ký kiểm toán.\n"
        "3. Triển khai thuật toán tối ưu vị trí đỗ xe tự động dựa trên vị trí địa lý của xe tại thời điểm đặt chỗ."
    )

    # --- Create docs folder if not exist ---
    if not os.path.exists('docs'):
        os.makedirs('docs')

    # Save Document
    doc_path = 'docs/Bao_cao_Do_an_Smart_Parking.docx'
    doc.save(doc_path)
    print(f"Document created successfully at {doc_path}")

if __name__ == '__main__':
    create_document()
