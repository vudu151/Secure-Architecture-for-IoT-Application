#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo file báo cáo Bài 4 ESP32 Sensor - đầy đủ theo yêu cầu nộp bài:
 1. Code chương trình với chú giải
 2. Link bài mô phỏng Wokwi
 3. Ảnh chụp màn hình kết quả mô phỏng
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
BRAIN_DIR = "/home/duvx/.gemini/antigravity/brain/c37b6250-3242-4b54-ae76-31e05f55ca5a"
IMG_MAIN   = os.path.join(BRAIN_DIR, "wokwi_simulation_main_1781102748230.png")
IMG_CIRCUIT= os.path.join(BRAIN_DIR, "wokwi_lcd_display_1781102769798.png")
IMG_SERIAL = os.path.join(BRAIN_DIR, "wokwi_serial_monitor_1781102793205.png")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.0)

# ── Set default font ──────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ── Helpers ───────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def font_run(run, size=13, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def heading(text, level=1, color='1F3864', size=None):
    p = doc.add_heading('', level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    default_size = {1: 15, 2: 13, 3: 12}.get(level, 13)
    font_run(run, size=size or default_size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(text='', bold=False, italic=False, size=13,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, color=None,
         space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        font_run(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_run(p, text, bold=False, italic=False, color=None, size=13):
    run = p.add_run(text)
    font_run(run, size=size, bold=bold, italic=italic, color=color)
    return run

def code_block(text, font_size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F4F8')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return p

def add_table(rows_data, headers=None, col_widths=None):
    cols = len(rows_data[0])
    t = doc.add_table(rows=0, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    if headers:
        hrow = t.add_row()
        for i, h in enumerate(headers):
            cell = hrow.cells[i]
            set_cell_bg(cell, '2E75B6')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            font_run(r, size=12, bold=True, color='FFFFFF')
    for ri, row in enumerate(rows_data):
        drow = t.add_row()
        bg = 'DEEAF1' if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row):
            cell = drow.cells[i]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            font_run(r, size=12)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t

def add_image(path, width_cm=15, caption=None):
    try:
        doc.add_picture(path, width=Cm(width_cm))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            r = cp.add_run(caption)
            font_run(r, size=11, italic=True, color='595959')
    except Exception as e:
        p = para(f'[Ảnh: {os.path.basename(path)} — {e}]', italic=True, color='CC0000')

# ═════════════════════════════════════════════════════════════════
# TRANG BÌA
# ═════════════════════════════════════════════════════════════════
def center_para(text, size=13, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    font_run(r, size=size, bold=bold, italic=italic, color=color)
    return p

center_para('TRƯỜNG ĐẠI HỌC BÁCH KHOA HÀ NỘI', size=13, bold=True, space_after=2)
center_para('Khoa Công nghệ Thông tin và Truyền thông', size=12, space_after=20)
center_para('')
center_para('BÁO CÁO BÀI TẬP', size=18, bold=True, color='1F3864', space_after=8)
center_para('BÀI 4: LẬP TRÌNH ESP32 THU THẬP DỮ LIỆU SENSOR', size=15, bold=True, color='2E75B6', space_after=6)
center_para('Môn học: Kiến trúc và Bảo mật cho Ứng dụng IoT', size=13, italic=True, space_after=30)

# info box
info_t = doc.add_table(rows=3, cols=2)
info_t.style = 'Table Grid'
info_t.alignment = WD_TABLE_ALIGNMENT.CENTER
info_rows = [
    ('Họ và tên sinh viên:', ''),
    ('Mã số sinh viên (MSSV):', ''),
    ('Ngày thực hiện:', '10/06/2026'),
]
for i, (lbl, val) in enumerate(info_rows):
    c0, c1 = info_t.rows[i].cells[0], info_t.rows[i].cells[1]
    set_cell_bg(c0, 'DEEAF1')
    c0.width = Cm(7); c1.width = Cm(7)
    r0 = c0.paragraphs[0].add_run(lbl); font_run(r0, size=12, bold=True)
    r1 = c1.paragraphs[0].add_run(val); font_run(r1, size=12)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# PHẦN I — CODE CHƯƠNG TRÌNH
# ═════════════════════════════════════════════════════════════════
heading('I. CODE CHƯƠNG TRÌNH VỚI CÁC CHÚ GIẢI CẦN THIẾT', level=1)

para('Chương trình được viết bằng C++ trên nền tảng Arduino cho ESP32, '
     'sử dụng hai thư viện: DHTesp (đọc cảm biến DHT22) và LiquidCrystal_I2C (điều khiển LCD I2C). '
     'Kỹ thuật millis() được áp dụng để đọc DHT22 định kỳ 2 giây mà không làm '
     'gián đoạn việc kiểm tra tín hiệu PIR liên tục.')

para('')

full_code = """\
/**
 * ================================================================
 * Bai 4: Lap trinh ESP32 thu thap du lieu sensor
 * Mon hoc: Kien truc va Bao mat cho Ung dung IoT
 * ================================================================
 * Chuc nang chinh:
 *  1. Doc nhiet do, do am tu cam bien DHT22 moi 2 giay
 *  2. Hien thi du lieu len LCD 16x2 qua giao tiep I2C
 *  3. Phat hien chuyen dong qua PIR -> bat/tat LED
 *  4. Gui log du lieu, thong bao qua Serial port (115200 baud)
 *
 * So do ket noi:
 *  DHT22  DATA  -> GPIO 15
 *  LCD    SDA   -> GPIO 21  |  SCL -> GPIO 22  (I2C 0x27)
 *  PIR    OUT   -> GPIO 12
 *  LED (+)      -> GPIO 32  qua dien tro 1kOhm -> GND
 * ================================================================
 */

#include "DHTesp.h"            // Thu vien giao tiep cam bien DHT cho ESP32
#include "LiquidCrystal_I2C.h" // Thu vien dieu khien LCD 16x2 qua I2C

// ----------------------------------------------------------------
// DINH NGHIA CHAN KET NOI (PIN definitions)
// ----------------------------------------------------------------
#define DHT_PIN           15   // Chan DATA cua DHT22 ket noi GPIO 15
#define PIR_PIN           12   // Chan OUT  cua PIR   ket noi GPIO 12
#define LED_PIN           32   // Chan dieu khien LED ket noi GPIO 32

// Chu ky doc cam bien DHT22: 2000ms = 2 giay
#define READ_INTERVAL_MS  2000

// ----------------------------------------------------------------
// KHOI TAO DOI TUONG
// ----------------------------------------------------------------
DHTesp dhtSensor;  // Doi tuong giao tiep voi cam bien DHT22

// Khoi tao LCD I2C:
//   Tham so: dia chi I2C = 0x27, 16 cot, 2 hang
LiquidCrystal_I2C lcd(0x27, 16, 2);

// ----------------------------------------------------------------
// BIEN TOAN CUC
// ----------------------------------------------------------------
unsigned long lastReadTime = 0; // Luu thoi diem doc DHT22 lan truoc (ms)
float currentTemp  = 0.0;       // Nhiet do hien tai (do C)
float currentHumid = 0.0;       // Do am hien tai (%)
bool  motionDetected = false;   // Trang thai: dang co chuyen dong hay khong

// ================================================================
// SETUP() - Chay 1 lan khi ESP32 khoi dong
// ================================================================
void setup() {
  // [1] Khoi dong cong Serial de in log ra may tinh
  Serial.begin(115200);
  Serial.println("==========================================");
  Serial.println("  ESP32 - Thu thap du lieu Sensor        ");
  Serial.println("==========================================");

  // [2] Cau hinh chan GPIO
  pinMode(LED_PIN, OUTPUT);    // GPIO 32: dau ra - dieu khien LED
  pinMode(PIR_PIN, INPUT);     // GPIO 12: dau vao - nhan tin hieu PIR
  digitalWrite(LED_PIN, LOW);  // Tat LED khi moi khoi dong

  // [3] Khoi dong cam bien DHT22 tren GPIO 15, kieu DHT22
  dhtSensor.setup(DHT_PIN, DHTesp::DHT22);
  Serial.println("[DHT22] Cam bien nhiet do/do am da san sang.");

  // [4] Khoi dong man hinh LCD
  lcd.init();      // Khoi tao LCD
  lcd.backlight(); // Bat den nen LCD (khong co backlight man hinh toi)

  // Hien thi man hinh chao trong 2 giay
  lcd.setCursor(0, 0);
  lcd.print("  ESP32 Sensor  ");
  lcd.setCursor(0, 1);
  lcd.print("  Initializing..");
  Serial.println("[LCD] Man hinh da san sang.");

  delay(2000);  // Cho 2 giay de man hinh chao duoc hien thi
  lcd.clear();  // Xoa man hinh truoc khi vao vong lap chinh

  Serial.println("[SYSTEM] He thong san sang!");
  Serial.println("------------------------------------------");
}

// ================================================================
// readDHTSensor() - Doc du lieu tu cam bien DHT22
// Luu ket qua vao bien toan cuc currentTemp va currentHumid
// ================================================================
void readDHTSensor() {
  // Lay du lieu nhiet do va do am cung luc
  TempAndHumidity data = dhtSensor.getTempAndHumidity();

  // isnan() kiem tra gia tri co hop le khong (NaN = Not a Number = loi)
  if (!isnan(data.temperature) && !isnan(data.humidity)) {
    currentTemp  = data.temperature; // Luu nhiet do hop le
    currentHumid = data.humidity;    // Luu do am hop le

    // In ket qua ra Serial Monitor
    Serial.println("[DHT22] Nhiet do: " + String(currentTemp, 1) + " C");
    Serial.println("[DHT22] Do am:    " + String(currentHumid, 1) + " %");
  } else {
    // Truong hop cam bien tra ve gia tri loi
    Serial.println("[DHT22] LOI: Khong doc duoc du lieu cam bien!");
  }
}

// ================================================================
// updateLCD() - Cap nhat hien thi len man hinh LCD
// Dong 0: nhiet do | Dong 1: do am
// ================================================================
void updateLCD() {
  // --- Dong 0: Hien thi nhiet do ---
  lcd.setCursor(0, 0);          // Di chuyen con tro den cot 0, hang 0
  lcd.print("Temp: ");          // In nhan "Temp:"
  lcd.print(String(currentTemp, 1)); // In gia tri nhiet do 1 chu so thap phan
  lcd.print((char)223);         // In ky tu do "°" (ma ASCII 223)
  lcd.print("C   ");            // In don vi va xoa ky tu cu (neu co)

  // --- Dong 1: Hien thi do am ---
  lcd.setCursor(0, 1);          // Di chuyen con tro den cot 0, hang 1
  lcd.print("Humid:");          // In nhan "Humid:"
  lcd.print(String(currentHumid, 1)); // In gia tri do am 1 chu so thap phan
  lcd.print("%   ");            // In don vi va xoa ky tu cu (neu co)
}

// ================================================================
// handlePIRSensor() - Xu ly tin hieu PIR va dieu khien LED
// HIGH = phat hien chuyen dong -> bat LED
// LOW  = khong co chuyen dong  -> tat LED
// ================================================================
void handlePIRSensor() {
  int pirValue = digitalRead(PIR_PIN); // Doc muc dien ap tai GPIO 12

  if (pirValue == HIGH) {
    // --- TRUONG HOP: Phat hien co chuyen dong ---
    if (!motionDetected) {
      // Chi in log khi trang thai MOI thay doi (tu OFF -> ON)
      // Tranh viec in qua nhieu dong log lap lai
      Serial.println("[PIR] *** Phat hien chuyen dong! LED ON ***");
      motionDetected = true;  // Cap nhat trang thai
    }
    digitalWrite(LED_PIN, HIGH); // Bat LED (muc dien ap cao = 3.3V)

  } else {
    // --- TRUONG HOP: Khong co chuyen dong ---
    if (motionDetected) {
      // Chi in log khi trang thai MOI thay doi (tu ON -> OFF)
      Serial.println("[PIR] Khong co chuyen dong. LED OFF.");
      motionDetected = false; // Cap nhat trang thai
    }
    digitalWrite(LED_PIN, LOW);  // Tat LED (muc dien ap thap = 0V)
  }
}

// ================================================================
// LOOP() - Vong lap chinh, chay lien tuc
// ================================================================
void loop() {
  // Lay thoi gian hien tai tinh tu luc ESP32 khoi dong (ms)
  unsigned long currentTime = millis();

  // ----------------------------------------------------------------
  // NHIEM VU 1: Doc DHT22 va cap nhat LCD moi 2 giay
  // Dung millis() (NON-BLOCKING) thay vi delay(2000):
  //   - delay(2000) se lam DONG BANG chuong trinh 2 giay
  //   - millis() cho phep xu ly PIR DONG THOI trong thoi gian cho
  // ----------------------------------------------------------------
  if (currentTime - lastReadTime >= READ_INTERVAL_MS) {
    lastReadTime = currentTime; // Cap nhat moc thoi gian

    readDHTSensor(); // Doc du lieu nhiet do / do am
    updateLCD();     // Cap nhat man hinh LCD
    Serial.println("------------------------------------------");
  }

  // ----------------------------------------------------------------
  // NHIEM VU 2: Xu ly PIR + dieu khien LED (LIEN TUC)
  // Duoc goi moi ~100ms, khong bi chan boi delay(2000)
  // ----------------------------------------------------------------
  handlePIRSensor();

  // Delay nho 100ms de:
  //   1. Giam tan so kiem tra PIR (tranh nhieu/false trigger)
  //   2. Tranh ESP32 chay vong lap qua nhanh (lam nong chip)
  delay(100);
}"""

code_block(full_code)

# ── Thư viện dùng ──────────────────────────────────────────────
doc.add_paragraph()
heading('Thư viện sử dụng (platformio.ini)', level=2, color='2E75B6', size=12)
code_block("""\
[env:esp32doit-devkit-v1]
platform  = espressif32
board     = esp32doit-devkit-v1
framework = arduino
lib_deps  =
    marcoschwartz/LiquidCrystal_I2C@^1.1.4
    beegee-tokyo/DHT sensor library for ESPx@^1.19.0""", font_size=10)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# PHẦN II — SƠ ĐỒ KẾT NỐI
# ═════════════════════════════════════════════════════════════════
heading('II. SƠ ĐỒ GHÉP NỐI ESP32 VÀ CÁC SENSOR', level=1)

add_table(
    rows_data=[
        ['DHT22 → DATA', 'ESP32 GPIO 15', 'Xanh lá'],
        ['DHT22 → VCC',  'ESP32 3.3V',   'Đỏ'],
        ['DHT22 → GND',  'ESP32 GND',    'Đen'],
        ['LCD I2C → SDA','ESP32 GPIO 21 (SDA)', 'Xanh lá'],
        ['LCD I2C → SCL','ESP32 GPIO 22 (SCL)', 'Xanh dương'],
        ['LCD I2C → VCC','ESP32 VCC (5V)',      'Đỏ'],
        ['LCD I2C → GND','ESP32 GND',           'Đen'],
        ['PIR → OUT',    'ESP32 GPIO 12',       'Xanh lá'],
        ['PIR → VCC',    'ESP32 3.3V',          'Đỏ'],
        ['PIR → GND',    'ESP32 GND',           'Đen'],
        ['LED Anode (+)','ESP32 GPIO 32',       'Xanh lá'],
        ['LED Cathode (−) → R(1kΩ)', 'ESP32 GND', 'Đen'],
    ],
    headers=['Từ (From)', 'Đến (To)', 'Màu dây'],
    col_widths=[5.5, 5.5, 3.5]
)

doc.add_paragraph()
para('Sơ đồ kết nối logic:')
code_block("""\
┌──────────────────────────────────────────────────────────┐
│                   ESP32 DevKit C v4                      │
│                                                          │
│  GPIO 15 ──────────────────────────► DHT22  (DATA)      │
│  GPIO 21 (SDA) ────────────────────► LCD I2C (SDA)      │
│  GPIO 22 (SCL) ────────────────────► LCD I2C (SCL)      │
│  GPIO 12 ◄──────────────────────────  PIR Motion (OUT)  │
│  GPIO 32 ──────── [R 1kΩ] ──────────► LED Anode (+)     │
│                                                          │
│  3.3V ──────────────────────────────► DHT22 VCC         │
│  3.3V ──────────────────────────────► PIR VCC           │
│  VCC  ──────────────────────────────► LCD VCC           │
│  GND  ──────────────────────────────► Tất cả GND        │
└──────────────────────────────────────────────────────────┘""")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# PHẦN III — LINK WOKWI
# ═════════════════════════════════════════════════════════════════
heading('III. LINK BÀI MÔ PHỎNG TRÊN WOKWI', level=1)

para('Bài mô phỏng được thiết kế và chạy thử trên nền tảng Wokwi '
     '(https://wokwi.com/) — công cụ mô phỏng phần cứng nhúng trực tuyến '
     'hỗ trợ ESP32, Arduino và nhiều vi điều khiển khác.')

doc.add_paragraph()

link_t = doc.add_table(rows=2, cols=2)
link_t.style = 'Table Grid'
link_t.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_bg(link_t.rows[0].cells[0], '1F3864')
set_cell_bg(link_t.rows[0].cells[1], '1F3864')
h0 = link_t.rows[0].cells[0].paragraphs[0].add_run('Thông tin')
h1 = link_t.rows[0].cells[1].paragraphs[0].add_run('Nội dung')
font_run(h0, size=12, bold=True, color='FFFFFF')
font_run(h1, size=12, bold=True, color='FFFFFF')

# Row 1
set_cell_bg(link_t.rows[1].cells[0], 'DEEAF1')
link_t.rows[1].cells[0].width = Cm(4.5)
link_t.rows[1].cells[1].width = Cm(10)
r0 = link_t.rows[1].cells[0].paragraphs[0].add_run('Link mô phỏng Wokwi')
font_run(r0, size=12, bold=True)
r1 = link_t.rows[1].cells[1].paragraphs[0].add_run('https://wokwi.com/projects/431499073455973377')
font_run(r1, size=12, color='2E75B6', italic=False)

doc.add_paragraph()
para('Dự án Wokwi bao gồm: mạch điện ESP32 với DHT22, LCD 16×2 I2C, '
     'PIR motion sensor và LED đỏ; code chương trình và cấu hình thư viện đầy đủ.')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# PHẦN IV — ẢNH CHỤP MÀN HÌNH
# ═════════════════════════════════════════════════════════════════
heading('IV. ẢNH CHỤP MÀN HÌNH KẾT QUẢ MÔ PHỎNG', level=1)

# ── Ảnh 1: Toàn bộ giao diện
heading('4.1. Giao diện tổng quan mô phỏng Wokwi', level=2, color='2E75B6', size=12)
para('Hình dưới hiển thị toàn bộ giao diện Wokwi khi mô phỏng đang chạy, '
     'bao gồm code editor (trái), sơ đồ mạch (phải trên) và Serial Monitor (phải dưới):')
doc.add_paragraph()
add_image(IMG_MAIN, width_cm=14.5,
          caption='Hình 1: Giao diện Wokwi — Simulation đang chạy, LED ON khi phát hiện chuyển động')

doc.add_paragraph()

# ── Ảnh 2: Mạch + LCD
heading('4.2. Sơ đồ mạch và kết quả hiển thị LCD', level=2, color='2E75B6', size=12)
para('Màn hình LCD 16×2 hiển thị nhiệt độ 37.0°C và độ ẩm 69.5% đọc từ cảm biến DHT22. '
     'LED đỏ đang bật sáng do cảm biến PIR phát hiện có chuyển động:')
doc.add_paragraph()
add_image(IMG_CIRCUIT, width_cm=12,
          caption='Hình 2: Sơ đồ mạch — LCD hiển thị Temp: 37.0°C / Humid:69.5%, LED ON (phát hiện chuyển động)')

doc.add_page_break()

# ── Ảnh 3: Serial Monitor
heading('4.3. Kết quả log trên Serial Monitor', level=2, color='2E75B6', size=12)
para('Serial Monitor ở tốc độ 115200 baud hiển thị log chương trình: '
     'dữ liệu nhiệt độ/độ ẩm cập nhật mỗi 2 giây, '
     'thông báo trạng thái PIR và LED theo từng chu kỳ:')
doc.add_paragraph()
add_image(IMG_SERIAL, width_cm=11,
          caption='Hình 3: Serial Monitor — Log dữ liệu DHT22 và trạng thái PIR/LED')

doc.add_paragraph()

# ── Mô tả kết quả ──────────────────────────────────────────────
heading('4.4. Mô tả kết quả', level=2, color='2E75B6', size=12)
add_table(
    rows_data=[
        ['LCD dòng 0', 'Temp: 37.0°C', 'Nhiệt độ đọc từ DHT22'],
        ['LCD dòng 1', 'Humid:69.5%',  'Độ ẩm đọc từ DHT22'],
        ['LED đỏ', 'ON (sáng)', 'PIR phát hiện chuyển động'],
        ['Serial log DHT22', '[DHT22] Nhiet do: 37.0 C', 'Cập nhật mỗi 2 giây'],
        ['Serial log PIR ON', '[PIR] *** Phat hien chuyen dong! LED ON ***', 'Khi PIR = HIGH'],
        ['Serial log PIR OFF','[PIR] Khong co chuyen dong. LED OFF.', 'Khi PIR = LOW'],
    ],
    headers=['Thành phần', 'Giá trị', 'Mô tả'],
    col_widths=[4, 6.5, 5]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════
# PHẦN V — KẾT LUẬN
# ═════════════════════════════════════════════════════════════════
heading('V. KẾT LUẬN', level=1)
para('Bài thực hành đã lập trình thành công hệ thống ESP32 thu thập dữ liệu '
     'từ đa cảm biến với các chức năng đầy đủ theo yêu cầu:')

conclusions = [
    'Giao tiếp DHT22 qua giao thức 1-Wire: đọc nhiệt độ, độ ẩm chính xác mỗi 2 giây.',
    'Hiển thị LCD 16×2 qua I2C: dữ liệu cảm biến được cập nhật theo thời gian thực.',
    'Điều khiển LED theo tín hiệu PIR Motion: phát hiện chuyển động và phản hồi ngay lập tức.',
    'Ghi log Serial 115200 baud: theo dõi trạng thái hệ thống đầy đủ, rõ ràng.',
    'Áp dụng kỹ thuật non-blocking (millis()): thực thi đa nhiệm hiệu quả trên vi điều khiển đơn nhân.',
    'Mô phỏng thành công trên Wokwi: kết quả đúng với yêu cầu đề bài.',
]
for c in conclusions:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, c, size=13)

# ─── Save ────────────────────────────────────────────────────────
out = os.path.join(BASE_DIR, "Bao_cao_Bai4_ESP32_Sensor.docx")
doc.save(out)
print(f"✅  Saved: {out}")
